from __future__ import annotations

import os
import re
import shutil
import subprocess
import tempfile
from dataclasses import dataclass
from typing import Optional, Tuple

from docx import Document  # python-docx
from PyPDF2 import PdfReader


@dataclass
class ExtractedText:
    filename: str
    text: str
    detected_type: str


def _clean_text(s: str) -> str:
    if not s:
        return ""
    s = s.replace("\x00", " ")
    s = re.sub(r"[ \t]+", " ", s)
    s = re.sub(r"\n{3,}", "\n\n", s)
    return s.strip()


def _read_docx(path: str) -> str:
    doc = Document(path)
    parts = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)
    # Tables (often important in design docs)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join((cell.text or "").strip() for cell in row.cells)
            row_text = row_text.strip()
            if row_text and row_text != "|":
                parts.append(row_text)
    return _clean_text("\n".join(parts))


def _read_pdf(path: str) -> str:
    reader = PdfReader(path)
    parts = []
    for page in reader.pages:
        try:
            txt = page.extract_text() or ""
        except Exception:
            txt = ""
        txt = txt.strip()
        if txt:
            parts.append(txt)
    return _clean_text("\n\n".join(parts))


def _read_txt(path: str) -> str:
    with open(path, "rb") as f:
        raw = f.read()
    # Try utf-8 first, fallback to latin-1
    try:
        text = raw.decode("utf-8")
    except UnicodeDecodeError:
        text = raw.decode("latin-1", errors="ignore")
    return _clean_text(text)


def _soffice_convert_to_docx(input_path: str, out_dir: str) -> Optional[str]:
    """
    Converts legacy .doc to .docx using LibreOffice headless (soffice).
    Returns the converted .docx path if successful.
    """
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        return None

    # LibreOffice will output same basename but .docx
    try:
        subprocess.run(
            [
                soffice,
                "--headless",
                "--convert-to",
                "docx",
                "--outdir",
                out_dir,
                input_path,
            ],
            check=True,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            timeout=60,
        )
    except Exception:
        return None

    base = os.path.splitext(os.path.basename(input_path))[0]
    converted = os.path.join(out_dir, f"{base}.docx")
    return converted if os.path.exists(converted) else None


def extract_text_from_upload(
    filename: str,
    file_bytes: bytes,
) -> ExtractedText:
    """
    Supports: .docx, .doc, .pdf, .txt, .md
    For .doc, attempts soffice conversion -> docx.
    """
    ext = (os.path.splitext(filename)[1] or "").lower().strip(".")
    detected = ext or "unknown"

    with tempfile.TemporaryDirectory() as td:
        in_path = os.path.join(td, filename)
        with open(in_path, "wb") as f:
            f.write(file_bytes)

        if ext in ("docx",):
            return ExtractedText(filename, _read_docx(in_path), "docx")

        if ext in ("pdf",):
            return ExtractedText(filename, _read_pdf(in_path), "pdf")

        if ext in ("txt", "md"):
            return ExtractedText(filename, _read_txt(in_path), "text")

        if ext in ("doc",):
            converted = _soffice_convert_to_docx(in_path, td)
            if not converted:
                # best-effort: return empty with guidance
                return ExtractedText(
                    filename,
                    "",
                    "doc (conversion_failed)"
                )
            return ExtractedText(filename, _read_docx(converted), "doc->docx")

        # Unknown type
        return ExtractedText(filename, "", f"unsupported:{detected}")
